from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List, Dict, Optional
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import copy

app = FastAPI()

# Enable CORS for frontend communication
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class ChatMessage(BaseModel):
    role: str
    message: str

class ChatRequest(BaseModel):
    message: str
    conversation_history: List[ChatMessage]

class GenerateRequest(BaseModel):
    conversation_history: List[ChatMessage]

def extract_placeholders(doc: Document) -> List[Dict[str, str]]:
    """Extract placeholders from document in format [Placeholder Name]"""
    placeholders = []
    seen = set()
    
    # Pattern to match [Text] or \[Text\]
    pattern = r'\[([^\]]+)\]'
    
    # Check paragraphs
    for para in doc.paragraphs:
        matches = re.findall(pattern, para.text)
        for match in matches:
            if match not in seen:
                seen.add(match)
                placeholders.append({
                    'name': match,
                    'value': None,
                    'question': f"What should I fill in for [{match}]?"
                })
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(pattern, cell.text)
                for match in matches:
                    if match not in seen:
                        seen.add(match)
                        placeholders.append({
                            'name': match,
                            'value': None,
                            'question': f"What should I fill in for [{match}]?"
                        })
    
    return placeholders

def extract_values_from_conversation(conversation: List[ChatMessage], placeholders: List[Dict]) -> Dict[str, str]:
    """Use simple pattern matching to extract values from conversation"""
    values = {}
    
    # Create a mapping of lowercase placeholder names for matching
    placeholder_names = {p['name'].lower(): p['name'] for p in placeholders}
    
    for i, msg in enumerate(conversation):
        if msg.role == 'assistant':
            # Check if the assistant asked about a placeholder
            for placeholder_key, placeholder_name in placeholder_names.items():
                if placeholder_key in msg.message.lower() or f"[{placeholder_name}]" in msg.message:
                    # Get the next user message
                    if i + 1 < len(conversation) and conversation[i + 1].role == 'user':
                        values[placeholder_name] = conversation[i + 1].message
                        break
    
    return values

def replace_text_in_run(run, placeholder, value):
    """Replace placeholder in a run while preserving formatting"""
    if placeholder in run.text:
        run.text = run.text.replace(placeholder, value)

def replace_text_in_paragraph(paragraph, placeholder, value):
    """Replace placeholder in paragraph while preserving all formatting"""
    full_text = paragraph.text
    if placeholder not in full_text:
        return
    
    # If the placeholder is within a single run, simple replacement
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return
    
    # If placeholder spans multiple runs, we need to be more careful
    # Build a map of character positions to runs
    char_to_run = []
    for run in paragraph.runs:
        char_to_run.extend([run] * len(run.text))
    
    # Find all occurrences of the placeholder
    while placeholder in paragraph.text:
        full_text = paragraph.text
        start_idx = full_text.find(placeholder)
        if start_idx == -1:
            break
        
        end_idx = start_idx + len(placeholder)
        
        # Clear the text in all runs
        for run in paragraph.runs:
            run.text = ""
        
        # Rebuild the text with replacement
        current_pos = 0
        for i, run in enumerate(paragraph.runs):
            if current_pos < start_idx:
                # Before placeholder
                chars_to_add = min(len(full_text) - current_pos, start_idx - current_pos)
                run.text = full_text[current_pos:current_pos + chars_to_add]
                current_pos += chars_to_add
            elif current_pos < end_idx:
                # At placeholder - put replacement in first run that hits it
                if not any(value in r.text for r in paragraph.runs[:i]):
                    run.text = value
                current_pos = end_idx
            else:
                # After placeholder
                run.text = full_text[current_pos:]
                break
        break

def fill_document_preserve_formatting(doc: Document, values: Dict[str, str]) -> Document:
    """Fill placeholders while preserving ALL formatting"""
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        for placeholder, value in values.items():
            placeholder_str = f"[{placeholder}]"
            replace_text_in_paragraph(paragraph, placeholder_str, value)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in values.items():
                        placeholder_str = f"[{placeholder}]"
                        replace_text_in_paragraph(paragraph, placeholder_str, value)
    
    # Process headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for placeholder, value in values.items():
                placeholder_str = f"[{placeholder}]"
                replace_text_in_paragraph(paragraph, placeholder_str, value)
        
        # Process footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            for placeholder, value in values.items():
                placeholder_str = f"[{placeholder}]"
                replace_text_in_paragraph(paragraph, placeholder_str, value)
    
    return doc

# Global storage for uploaded document (as bytes to preserve everything)
current_doc_bytes = None
current_placeholders = []

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a legal document"""
    global current_doc_bytes, current_placeholders
    
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        # Read and store the original file bytes
        contents = await file.read()
        current_doc_bytes = contents
        
        # Parse document for placeholder extraction
        doc = Document(BytesIO(contents))
        
        # Extract placeholders
        placeholders = extract_placeholders(doc)
        current_placeholders = placeholders
        
        return {
            "filename": file.filename,
            "placeholders": placeholders,
            "count": len(placeholders)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.post("/chat")
async def chat(request: ChatRequest):
    """Handle conversation to fill placeholders"""
    global current_placeholders
    
    try:
        # Extract values from conversation so far
        values = extract_values_from_conversation(request.conversation_history, current_placeholders)
        
        # Find next unfilled placeholder
        next_placeholder = None
        for p in current_placeholders:
            if p['name'] not in values:
                next_placeholder = p
                break
        
        # Generate response
        if next_placeholder:
            response = f"Thank you! Now, {next_placeholder['question']}"
            all_filled = False
        else:
            response = "Great! I have all the information I need. Your document is ready to download."
            all_filled = True
        
        return {
            "response": response,
            "all_filled": all_filled,
            "filled_count": len(values),
            "total_count": len(current_placeholders)
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error in chat: {str(e)}")

@app.post("/generate")
async def generate_document(request: GenerateRequest):
    """Generate the completed document while preserving all formatting"""
    global current_doc_bytes, current_placeholders
    
    if current_doc_bytes is None:
        raise HTTPException(status_code=400, detail="No document uploaded")
    
    try:
        # Load the original document from stored bytes
        doc = Document(BytesIO(current_doc_bytes))
        
        # Extract all values from conversation
        values = extract_values_from_conversation(request.conversation_history, current_placeholders)
        
        # Fill the document while preserving formatting
        filled_doc = fill_document_preserve_formatting(doc, values)
        
        # Save to BytesIO
        doc_io = BytesIO()
        filled_doc.save(doc_io)
        doc_io.seek(0)
        
        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=completed_document.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")

@app.get("/")
async def root():
    return {"message": "Legal Document Assistant API", "status": "running"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)