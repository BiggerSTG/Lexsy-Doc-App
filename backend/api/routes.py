from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from io import BytesIO
import mammoth

from backend.core.models import ChatRequest, GenerateRequest
from backend.core.utils import (
    extract_placeholders,
    extract_values_from_conversation,
    fill_document_preserve_formatting,
)
from backend.storage import set_document, get_document, set_placeholders, get_placeholders
from docx import Document


router = APIRouter()


@router.post('/upload')
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a .docx legal document"""
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")

    try:
        contents = await file.read()
        set_document(contents)

        doc = Document(BytesIO(contents))
        placeholders = extract_placeholders(doc)
        set_placeholders(placeholders)

        return {
            "filename": file.filename,
            "placeholders": placeholders,
            "count": len(placeholders),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")


@router.post('/chat')
async def chat(request: ChatRequest):
    """Handle conversation to fill placeholders"""
    try:
        current_placeholders = get_placeholders()
        values = extract_values_from_conversation(request.conversation_history, current_placeholders)

        next_placeholder = None
        for p in current_placeholders:
            if p['name'] not in values:
                next_placeholder = p
                break

        if next_placeholder:
            response = f"Thank you! Now, {next_placeholder['question']}"
            all_filled = False
        else:
            response = "Great! I have all the information I need. Your document is ready to download."
            all_filled = True

        return {
            "response": response,
            "all_filled": all_filled,
            "filled_count": len(values),
            "total_count": len(current_placeholders),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error in chat: {str(e)}")


@router.post('/generate')
async def generate_document(request: GenerateRequest):
    """Generate the completed document while preserving formatting"""
    current_doc_bytes = get_document()
    current_placeholders = get_placeholders()
    if current_doc_bytes is None:
        raise HTTPException(status_code=400, detail="No document uploaded")

    try:
        doc = Document(BytesIO(current_doc_bytes))
        values = extract_values_from_conversation(request.conversation_history, current_placeholders)
        filled_doc = fill_document_preserve_formatting(doc, values)

        doc_io = BytesIO()
        filled_doc.save(doc_io)
        doc_io.seek(0)

        return StreamingResponse(
            doc_io,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=completed_document.docx"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating document: {str(e)}")


@router.post('/preview')
async def preview_document(request: GenerateRequest):
    current_doc_bytes = get_document()
    current_placeholders = get_placeholders()
    if current_doc_bytes is None:
        raise HTTPException(status_code=400, detail="No document uploaded")

    try:
        doc = Document(BytesIO(current_doc_bytes))
        values = extract_values_from_conversation(request.conversation_history, current_placeholders)
        filled_doc = fill_document_preserve_formatting(doc, values)

        preview_text = []
        for para in filled_doc.paragraphs:
            if para.text.strip():
                preview_text.append(para.text)

        for table in filled_doc.tables:
            preview_text.append("\n[TABLE]")
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    preview_text.append(row_text)
            preview_text.append("[END TABLE]\n")

        preview_content = "\n\n".join(preview_text)

        return {
            "preview": preview_content,
            "filled_count": len(values),
            "total_count": len(current_placeholders),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating preview: {str(e)}")


@router.post('/preview_html')
async def preview_document_html(request: GenerateRequest):
    current_doc_bytes = get_document()
    current_placeholders = get_placeholders()
    if current_doc_bytes is None:
        raise HTTPException(status_code=400, detail="No document uploaded")

    try:
        doc = Document(BytesIO(current_doc_bytes))
        values = extract_values_from_conversation(request.conversation_history, current_placeholders)
        filled_doc = fill_document_preserve_formatting(doc, values)

        doc_io = BytesIO()
        filled_doc.save(doc_io)
        doc_io.seek(0)

        result = mammoth.convert_to_html(doc_io)
        html = result.value

        return {
            "html": html,
            "filled_count": len(values),
            "total_count": len(current_placeholders),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating HTML preview: {str(e)}")


@router.get('/')
async def root():
    return {"message": "Legal Document Assistant API", "status": "running"}
