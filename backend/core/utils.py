import re
from typing import List, Dict
from docx import Document
from io import BytesIO


def extract_placeholders(doc: Document) -> List[Dict[str, str]]:
    """Extract placeholders from document in format [Placeholder Name]"""
    placeholders = []
    seen = set()
    
    # Pattern to match [Text] or \[Text\]
    pattern = r'\[([^\]]+)\]'
    
    # Check paragraphs
    for para in doc.paragraphs:
        text = para.text
        matches = re.finditer(pattern, text)
        for match in matches:
            placeholder = match.group(1)
            
            # Skip if it's just underscores or blank
            if placeholder.replace('_', '').strip() == '':
                # Look for quoted text after the placeholder
                after_text = text[match.end():]
                quote_match = re.search(r'"([^"]+)"', after_text[:100])  # Search in next 100 chars
                
                if not quote_match:
                    # Look for quoted text before the placeholder
                    before_text = text[:match.start()]
                    # Find the last occurrence of quoted text before placeholder
                    quote_matches = list(re.finditer(r'"([^"]+)"', before_text))
                    if quote_matches:
                        quote_match = quote_matches[-1]  # Get the closest one
                
                if quote_match:
                    descriptive_name = quote_match.group(1)
                    placeholder = f"{descriptive_name} in $"
                else:
                    placeholder = "Amount in $"
            
            if placeholder not in seen and not placeholder.replace('_', '').strip() == '':
                seen.add(placeholder)
                placeholders.append({
                    'name': placeholder,
                    'value': None,
                    'question': f"What should I fill in for [{placeholder}]?"
                })
    
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                matches = re.finditer(pattern, text)
                for match in matches:
                    placeholder = match.group(1)
                    
                    # Skip if it's just underscores
                    if placeholder.replace('_', '').strip() == '':
                        # Look for quoted text after
                        after_text = text[match.end():]
                        quote_match = re.search(r'"([^"]+)"', after_text[:100])
                        
                        if not quote_match:
                            # Look for quoted text before
                            before_text = text[:match.start()]
                            quote_matches = list(re.finditer(r'"([^"]+)"', before_text))
                            if quote_matches:
                                quote_match = quote_matches[-1]
                        
                        if quote_match:
                            descriptive_name = quote_match.group(1)
                            placeholder = f"{descriptive_name} in $"
                        else:
                            placeholder = "Amount in $"
                    
                    if placeholder not in seen and not placeholder.replace('_', '').strip() == '':
                        seen.add(placeholder)
                        placeholders.append({
                            'name': placeholder,
                            'value': None,
                            'question': f"What should I fill in for [{placeholder}]?"
                        })
    
    return placeholders


def extract_values_from_conversation(conversation: List[dict], placeholders: List[Dict]) -> Dict[str, str]:
    """Use simple pattern matching to extract values from conversation"""
    values = {}
    placeholder_names = {p['name'].lower(): p['name'] for p in placeholders}

    for i, msg in enumerate(conversation):
        if getattr(msg, 'role', None) == 'assistant' or (isinstance(msg, dict) and msg.get('role') == 'assistant'):
            # Check if the assistant asked about a placeholder
            msg_text = msg.message if hasattr(msg, 'message') else msg.get('message', '')
            for placeholder_key, placeholder_name in placeholder_names.items():
                if placeholder_key in msg_text.lower() or f"[{placeholder_name}]" in msg_text:
                    # Get the next user message
                    if i + 1 < len(conversation):
                        next_msg = conversation[i + 1]
                        next_text = next_msg.message if hasattr(next_msg, 'message') else next_msg.get('message', '')
                        values[placeholder_name] = next_text
                        break

    return values


def replace_text_in_paragraph(paragraph, placeholder, value):
    """Replace placeholder in paragraph while preserving run formatting where possible"""
    full_text = paragraph.text
    if placeholder not in full_text:
        return

    # If the placeholder is within a single run, simple replacement
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return

    # If placeholder spans multiple runs, rebuild runs carefully
    while placeholder in paragraph.text:
        full_text = paragraph.text
        start_idx = full_text.find(placeholder)
        if start_idx == -1:
            break
        end_idx = start_idx + len(placeholder)

        # Clear text in runs
        for run in paragraph.runs:
            run.text = ""

        # Rebuild text with replacement
        current_pos = 0
        for i, run in enumerate(paragraph.runs):
            if current_pos < start_idx:
                chars_to_add = min(len(full_text) - current_pos, start_idx - current_pos)
                run.text = full_text[current_pos:current_pos + chars_to_add]
                current_pos += chars_to_add
            elif current_pos < end_idx:
                # Put replacement in the first run that intersects
                if not any(value in r.text for r in paragraph.runs[:i]):
                    run.text = value
                current_pos = end_idx
            else:
                run.text = full_text[current_pos:]
                break
        break


def fill_document_preserve_formatting(doc: Document, values: Dict[str, str]) -> Document:
    """Fill placeholders while preserving ALL formatting"""
    
    # Create a mapping that includes underscore patterns
    replacement_map = {}
    for placeholder, value in values.items():
        # Standard bracketed placeholder
        replacement_map[f"[{placeholder}]"] = value
        
        # If it ends with "in $", also map the underscore pattern
        if placeholder.endswith(" in $"):
            # Map underscore patterns with $ prefix
            replacement_map[f"$[_____________]"] = f"${value}"
            replacement_map[f"[_____________]"] = value
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        for placeholder_str, value in replacement_map.items():
            replace_text_in_paragraph(paragraph, placeholder_str, value)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder_str, value in replacement_map.items():
                        replace_text_in_paragraph(paragraph, placeholder_str, value)
    
    # Process headers
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for placeholder_str, value in replacement_map.items():
                replace_text_in_paragraph(paragraph, placeholder_str, value)
        
        # Process footer
        footer = section.footer
        for paragraph in footer.paragraphs:
            for placeholder_str, value in replacement_map.items():
                replace_text_in_paragraph(paragraph, placeholder_str, value)
    
    return doc
